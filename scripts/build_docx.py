# -*- coding: utf-8 -*-
"""Word(.docx) 파일 생성: data/days/*.json → docs/files/YYYY-MM-DD_OPIc훈련.docx

이동 중 암기용이므로 스크립트·단어 표는 12pt 큰 글자로.
주말(weekend) 타입은 스크립트가 없으므로 건너뜀.
"""
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "data"
FILES = ROOT / "docs" / "files"

ACCENT = RGBColor(0x63, 0x66, 0xF1)
MUTED = RGBColor(0x6B, 0x72, 0x80)


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def heading(doc, text, size=14):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = ACCENT
    return p


def para(doc, text, size=12, bold=False, color=None, en=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if en:
        r.font.name = "Georgia"
    return p


def build_docx(day: dict, out: Path) -> None:
    doc = Document()

    # 제목
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"OPIc AL 데일리 훈련 — {day['date']}")
    r.bold = True
    r.font.size = Pt(18)
    para(doc, f"Unit {day.get('unit')} · {day.get('unit_title', '')} — {day.get('topic', '')}",
         size=11, color=MUTED)

    # ① 질문
    heading(doc, "① 오늘의 질문")
    para(doc, day["question"]["english"], size=12, bold=True, en=True)
    para(doc, day["question"]["korean"], size=11, color=MUTED)

    # ② 모범답변
    heading(doc, "② AL 모범답변")
    for p_txt in day["model_answer"]["script"].split("\n\n"):
        if p_txt.strip():
            para(doc, p_txt.strip(), size=12, en=True)

    # ③ 단어 15개 표 (단어/뜻/예문 — 접어서 뜻 가리고 외우기 좋게)
    heading(doc, "③ 오늘의 단어 15 (핵심 10 + 확장 5)")
    para(doc, "※ 오른쪽 두 칸을 접거나 가리고 뜻을 떠올려 보세요. '↔'는 대체 관계입니다.",
         size=9, color=MUTED)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(("단어", "뜻", "오늘 예문")):
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
    for w in day["vocab"]:
        row = table.add_row().cells
        word = w["word"]
        if w.get("tier") == "extend" and w.get("linked_word"):
            word += f"\n↔ {w['linked_word']} 대체"
        for cell, text, size in ((row[0], word, 12), (row[1], w["meaning"], 12),
                                 (row[2], w["example"], 11)):
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(size)
            if cell is row[0]:
                run.bold = True

    # ④ 전략 해설
    heading(doc, "④ 전략 해설 (AL 포인트)")
    para(doc, "답변 뼈대: " + day["model_answer"]["structure_note"], size=10, color=MUTED)
    for pt in day["model_answer"]["al_points"]:
        para(doc, "✔ " + pt, size=11)

    # ⑤ 패턴
    heading(doc, "⑤ 오늘의 패턴")
    for p in day["patterns"]:
        para(doc, p["pattern"], size=12, bold=True, en=True)
        para(doc, f"{p['meaning']} / 응용: {p['extra_example']}", size=10, color=MUTED)

    # ⑥ 복습 단어
    if day.get("review_words"):
        heading(doc, "⑥ 복습 단어 (3일·7일 전)")
        for w in day["review_words"]:
            para(doc, f"{w['word']} — {w['new_example']}", size=11)

    # ⑦ 체크리스트
    heading(doc, "⑦ 셀프 체크리스트")
    for c in day["checklist"]:
        para(doc, "□ " + c, size=11)

    doc.save(out)


def main() -> None:
    FILES.mkdir(parents=True, exist_ok=True)
    built = 0
    for f in sorted((DATA / "days").glob("*.json")):
        day = load_json(f)
        if day.get("type") == "weekend":
            continue
        out = FILES / f"{day['date']}_OPIc훈련.docx"
        build_docx(day, out)
        built += 1
    print(f"[build_docx] {built}개 Word 파일 생성 → docs/files/")


if __name__ == "__main__":
    main()
